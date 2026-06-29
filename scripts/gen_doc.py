# -*- coding: utf-8 -*-
"""
Generate a Word document from a JSON config, based on the standard Yingdao RPA requirements template.

Usage:
    python gen_doc.py <config.json>

The config.json structure:
{
    "template_path": "C:/Users/13251/.workbuddy/skills/video-to-rpa-doc/references/影刀RPA需求文档模板.docx",
    "output_path": "path/to/output.docx",
    "frames_dir": "path/to/extracted/frames",
    "title": "Document title",
    "subtitle": "",                  // optional; if empty, subtitle is cleared
    "version": "Version 1.0",
    "date": "2026.06.28",            // or "auto"
    "basic_info": {
        "name": "Process name",
        "department": "Department",
        "description": "Scene description",
        "duration": "10秒",
        "daily_count": "500",
        "remarks": "Remarks"
    },
    "systems": [
        ["System name", "内部/外部", "Windows/Web/iOS", "Yes/No", "Yes/No", "无", "Remarks"],
        ...
    ],
    "steps": [
        {
            "name": "Step 1 name(action)",
            "desc": "Detailed description with line breaks",
            "supplement": "Supplementary notes",
            "img": "frame_00_5s.png",   // filename relative to frames_dir
            "img_w": 5                    // image width in cm
        },
        ...
    ]
}
"""
import os
import sys
import json
import shutil
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_cell_text(cell, text, size=Pt(9), bold=False):
    """Safely set cell text, supporting multi-line via \\n."""
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p = cell.paragraphs[0]
    for run in p.runs:
        run._element.getparent().remove(run._element)
    lines = text.split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            br_run = p.add_run()
            br_run.add_break()
        run = p.add_run(line)
        run.font.size = size
        run.font.bold = bold


def set_cell_shading(cell, fill_color):
    """Set cell background color (hex like 'D9E2F3')."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = tcPr.makeelement(qn('w:shd'))
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill_color)
    shd.set(qn('w:val'), 'clear')


def add_image_to_cell(cell, img_path, width_cm=5):
    """Clear cell, then add a centered image."""
    if not os.path.exists(img_path):
        print(f"  WARNING: image not found: {img_path}")
        return
    # Remove all existing paragraphs in the cell
    for p in cell.paragraphs:
        p._element.getparent().remove(p._element)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Cm(width_cm))


def remove_extra_rows(table, keep_rows):
    """Remove extra <w:tr> rows from a table, keeping only the first `keep_rows` rows."""
    tbl_xml = table._tbl
    all_trs = list(tbl_xml.findall(qn('w:tr')))
    rows_to_remove = all_trs[keep_rows:]
    for row in rows_to_remove:
        tbl_xml.remove(row)


# XML namespaces for WPS textbox handling
NS_WPS = '{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}'
NS_W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'


def _set_txbx_p_text(p_elem, text):
    """Set text of a w:p inside a textbox, preserving font formatting from the first run."""
    # Save rPr from the first original run (run-level formatting like w:rFonts, w:sz)
    orig_runs = list(p_elem.findall(f'{NS_W}r'))
    saved_rPr = None
    for orig_run in orig_runs:
        rPr = orig_run.find(f'{NS_W}rPr')
        if rPr is not None:
            saved_rPr = rPr
            break

    # Remove all original runs
    for run in orig_runs:
        p_elem.remove(run)

    # Reconstruct font info: prefer run-level rPr, fall back to pPr/rPr
    pPr = p_elem.find(f'{NS_W}pPr')
    merged_rPr = p_elem.makeelement(f'{NS_W}rPr')

    # Copy pPr-level rPr children (paragraph defaults)
    if pPr is not None:
        pPr_rPr = pPr.find(f'{NS_W}rPr')
        if pPr_rPr is not None:
            for child in list(pPr_rPr):
                tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                # Don't copy if run-level already has it
                if saved_rPr is not None and saved_rPr.find(f'{NS_W}{tag}') is not None:
                    continue
                merged_rPr.append(child)

    # Copy run-level rPr children (explicit font info, takes priority)
    if saved_rPr is not None:
        for child in list(saved_rPr):
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            # Remove existing pPr-level element with same tag
            existing = merged_rPr.find(f'{NS_W}{tag}')
            if existing is not None:
                merged_rPr.remove(existing)
            merged_rPr.append(child)

    r = p_elem.makeelement(f'{NS_W}r')
    if len(merged_rPr) > 0:
        r.append(merged_rPr)

    t = p_elem.makeelement(f'{NS_W}t')
    t.text = text
    t.set(XML_SPACE, 'preserve')
    r.append(t)
    p_elem.append(r)


def update_cover_textboxes(doc_element, title, subtitle, version, date_str):
    """
    Update floating textbox content on the cover page.
    The template has WPS textboxes for:
      - Tagline ("一款顶级易用的流程自动化产品") — DO NOT MODIFY
      - Title ("流程名称")
      - Version / Date
      - Subtitle ("流程简介/副标题") — if present
    WSP shapes without textboxes (logo, decorative line) are left untouched.
    """
    for wsp in list(doc_element.iter(f'{NS_WPS}wsp')):
        txbx = wsp.find(f'{NS_WPS}txbx')
        if txbx is None:
            continue
        txbxContent = txbx.find(f'{NS_W}txbxContent')
        if txbxContent is None:
            continue

        all_text = ''.join(txbxContent.itertext())

        # Tagline textbox: NEVER modify (eg. "一款顶级易用的流程自动化产品")
        if '流程自动化' in all_text or '顶级易用' in all_text:
            continue

        paragraphs = list(txbxContent.findall(f'{NS_W}p'))

        if '流程名称' in all_text and '流程简介' not in all_text and '副标题' not in all_text:
            _set_txbx_p_text(paragraphs[0], title)
            for p in paragraphs[1:]:
                _set_txbx_p_text(p, '')

        elif '流程简介' in all_text or '副标题' in all_text:
            _set_txbx_p_text(paragraphs[0], subtitle)
            for p in paragraphs[1:]:
                _set_txbx_p_text(p, '')

        elif 'Version' in all_text:
            if len(paragraphs) >= 1:
                _set_txbx_p_text(paragraphs[0], version)
            if len(paragraphs) >= 2:
                _set_txbx_p_text(paragraphs[1], date_str)
            for p in paragraphs[2:]:
                _set_txbx_p_text(p, '')


NS_W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

def _para_is_empty(p_elem):
    """Check if a w:p element has no real content (text, image, drawing)."""
    runs = p_elem.findall(f'{NS_W}r')
    if not runs:
        return True
    for r in runs:
        t = r.find(f'{NS_W}t')
        if t is not None and t.text and t.text.strip():
            return False
        if r.find(f'{NS_W}drawing') is not None:
            return False
    return True

def remove_trailing_blank_page(doc):
    """
    Remove trailing blank paragraphs including those with page breaks,
    and clean up lastRenderedPageBreak artifacts that may cause an extra
    blank page at the end of the document.
    """
    body = doc.element.body
    children = list(body)
    ns_w = NS_W

    # --- Pass 1: Remove all lastRenderedPageBreak elements globally ---
    for lrpb in body.iter(f'{ns_w}lastRenderedPageBreak'):
        parent = lrpb.getparent()  # the w:r parent
        if parent is not None:
            parent.remove(lrpb)

    # --- Pass 2: Remove trailing empty paragraphs from the end ---
    removed = 0
    while True:
        children = list(body)
        if not children:
            break
        last = children[-1]
        tag = last.tag.split('}')[-1] if '}' in last.tag else last.tag
        if tag == 'sectPr':
            # Last element is section properties — check the one before it
            if len(children) >= 2 and children[-2].tag.split('}')[-1] == 'p':
                if _para_is_empty(children[-2]):
                    body.remove(children[-2])
                    removed += 1
                    continue
            break
        elif tag == 'p' and _para_is_empty(last):
            body.remove(last)
            removed += 1
            continue
        else:
            break

    # --- Pass 3: Remove any w:br type="page" inside trailing empty paragraphs ---
    # (template may have page breaks between sections that end up as blanks)
    children = list(body)
    # Find the last non-sectPr element
    last_idx = None
    for i in range(len(children) - 1, -1, -1):
        tag = children[i].tag.split('}')[-1]
        if tag != 'sectPr':
            last_idx = i
            break
    if last_idx is not None and last_idx >= 0:
        # Check if any paragraphs after the last table contain page breaks
        # that now produce a blank page
        for i in range(last_idx, -1, -1):
            tag = children[i].tag.split('}')[-1]
            if tag == 'tbl':
                # Found last table — remove empty paragraphs between table and end
                for j in range(len(children) - 1, i, -1):
                    tj = children[j].tag.split('}')[-1]
                    if tj == 'p' and _para_is_empty(children[j]):
                        body.remove(children[j])
                        removed += 1
                    elif tj == 'sectPr':
                        continue
                    else:
                        break
                break

    if removed > 0:
        print(f'  Removed {removed} trailing blank paragraph(s)')


def generate_doc(config):
    template_path = config['template_path']
    output_path = config['output_path']
    frames_dir = config.get('frames_dir', '')

    # Copy template -> output
    shutil.copy2(template_path, output_path)
    doc = Document(output_path)

    # --- Cover page ---
    title = config.get('title', '')
    subtitle = config.get('subtitle', '')
    version = config.get('version', 'Version 1.0')
    date_str = config.get('date', '')
    if date_str == 'auto' or not date_str:
        date_str = datetime.now().strftime('%Y.%m.%d')

    update_cover_textboxes(doc.element, title, subtitle, version, date_str)

    # --- Table 0: Basic info ---
    t0 = doc.tables[0]
    bi = config.get('basic_info', {})
    set_cell_text(t0.rows[1].cells[0], bi.get('name', ''))
    set_cell_text(t0.rows[1].cells[1], bi.get('department', ''))
    set_cell_text(t0.rows[1].cells[2], bi.get('description', ''))
    set_cell_text(t0.rows[1].cells[3], bi.get('duration', ''))
    set_cell_text(t0.rows[1].cells[4], bi.get('daily_count', ''))
    set_cell_text(t0.rows[1].cells[5], bi.get('remarks', ''))

    # --- Table 1: Systems ---
    t1 = doc.tables[1]
    systems = config.get('systems', [])
    for ri, vals in enumerate(systems):
        for ci, val in enumerate(vals):
            if ci < len(t1.columns):
                set_cell_text(t1.rows[ri + 1].cells[ci], val)
    # Remove extra rows: keep header + len(systems) data rows
    remove_extra_rows(t1, 1 + len(systems))

    # --- Keep "详细步骤解析" heading; only trim excess blank paragraphs before it ---
    body = doc.element.body
    children = list(body)
    ns_w = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    para_to_remove = []
    tbl_index = None
    heading_index = None
    for i, el in enumerate(children):
        tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
        if tag == 'p':
            text = ''.join(el.itertext()).strip()
            if '详细步骤解析' in text:
                heading_index = i
        elif tag == 'tbl':
            tbl_index = i
            if heading_index is not None:
                break

    # Trim excess blank paragraphs between heading and table (keep 1 blank para for spacing)
    if tbl_index is not None and heading_index is not None:
        blank_count = 0
        for i in range(tbl_index - 1, heading_index, -1):
            el = children[i]
            tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
            if tag == 'p' and not ''.join(el.itertext()).strip():
                blank_count += 1
                if blank_count > 1:
                    para_to_remove.append(i)
            else:
                break

    for idx in sorted(set(para_to_remove), reverse=True):
        body.remove(children[idx])

    # --- Table 2: Steps (3 cols: 步骤名称及详细说明 | 截图 | 其他补充说明及资料) ---
    # Each step occupies 2 rows:
    #   Row 1 (step title):   col 0 = step name (light blue bg)
    #   Row 2 (detail):       col 0 = description, col 1 = image, col 2 = supplement
    t2 = doc.tables[2]

    # Fix table header: template column titles may have old names
    set_cell_text(t2.rows[0].cells[0], '步骤名称及详细说明', bold=True)
    set_cell_text(t2.rows[0].cells[1], '截图', bold=True)
    steps = config.get('steps', [])
    for idx, step in enumerate(steps):
        r_title  = idx * 2 + 1
        r_detail = r_title + 1

        if r_title >= len(t2.rows):
            print(f"  WARNING: step {idx+1} exceeds template rows, skipping")
            break

        # 步骤号行：col0=步骤名称，col1/col2留空，加浅蓝背景
        # 步骤名称直接使用 config 中的原始值，不自动加编号前缀
        step_name = step.get('name', '')
        set_cell_text(t2.rows[r_title].cells[0], step_name, bold=True)
        set_cell_text(t2.rows[r_title].cells[1], '')
        set_cell_text(t2.rows[r_title].cells[2], '')
        for cell in t2.rows[r_title].cells:
            set_cell_shading(cell, 'E3F2FD')  # 极浅蓝

        if r_detail < len(t2.rows):
            # 详细说明行：描述放在 col 0（不含步骤名称），截图 col1，补充说明 col2
            set_cell_text(t2.rows[r_detail].cells[0], step.get('desc', ''))
            # 截图放在 col 1：add_image_to_cell 会自动清空旧内容再插入图片
            img_file = step.get('img', '')
            if img_file:
                img_path = os.path.join(frames_dir, img_file) if frames_dir else img_file
                add_image_to_cell(t2.rows[r_detail].cells[1], img_path, step.get('img_w', 5))
            else:
                set_cell_text(t2.rows[r_detail].cells[1], '')
            # 补充说明放在 col 2
            set_cell_text(t2.rows[r_detail].cells[2], step.get('supplement', ''))

    # Remove extra rows: keep header + len(steps) * 2 rows
    remove_extra_rows(t2, 1 + len(steps) * 2)

    # --- Remove trailing blank pages ---
    remove_trailing_blank_page(doc)

    # --- Save ---
    doc.save(output_path)
    print(f'Done: {output_path}')
    return output_path


def regenerate_flowchart(config):
    """
    Generate flowchart from the same config and return paths.
    Call this after generate_doc() to also produce flowchart files.
    """
    import importlib.util
    script_dir = os.path.dirname(os.path.abspath(__file__))
    flowchart_script = os.path.join(script_dir, 'gen_flowchart.py')
    
    # Determine output directory from the Word output path
    word_output = config.get('output_path', '')
    output_dir = os.path.dirname(word_output) if word_output else os.getcwd()
    
    if os.path.exists(flowchart_script):
        spec = importlib.util.spec_from_file_location("gen_flowchart", flowchart_script)
        gen_flowchart = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(gen_flowchart)
        return gen_flowchart.generate_flowchart(config, output_dir)
    else:
        print(f"  WARNING: flowchart script not found at {flowchart_script}")
        return None


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python gen_doc.py <config.json>")
        sys.exit(1)
    with open(sys.argv[1], 'r', encoding='utf-8') as f:
        config = json.load(f)
    generate_doc(config)
    # Note: Mermaid flowchart generation has been replaced by ProcessOn.
    # Use processon-diagram-generator skill to generate the flowchart instead.
    # Keep regenerate_flowchart() as fallback for local Mermaid generation.
